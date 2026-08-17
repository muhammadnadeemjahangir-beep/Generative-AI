import os
import time
from pathlib import Path

import pandas as pd
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS


# --------------------------------------------------
# Load environment variables
# --------------------------------------------------
load_dotenv()


# --------------------------------------------------
# File loaders
# --------------------------------------------------
def load_txt(file_path):
    file_path = Path(file_path)

    text = file_path.read_text(
        encoding="utf-8",
        errors="ignore"
    )

    if not text.strip():
        return []

    return [
        Document(
            page_content=text.strip(),
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "txt"
            }
        )
    ]


def load_pdf(file_path):
    file_path = Path(file_path)
    docs = []

    reader = PdfReader(str(file_path))

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text()

        if text and text.strip():
            docs.append(
                Document(
                    page_content=text.strip(),
                    metadata={
                        "source": str(file_path),
                        "file_name": file_path.name,
                        "file_type": "pdf",
                        "page": page_number
                    }
                )
            )

    return docs


def load_docx(file_path):
    file_path = Path(file_path)
    doc = DocxDocument(str(file_path))

    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                if cell.text and cell.text.strip():
                    row_text.append(cell.text.strip())

            if row_text:
                text_parts.append(" | ".join(row_text))

    text = "\n".join(text_parts)

    if not text.strip():
        return []

    return [
        Document(
            page_content=text.strip(),
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "docx"
            }
        )
    ]


def load_csv(file_path):
    file_path = Path(file_path)

    try:
        df = pd.read_csv(file_path)
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, encoding="latin1")

    if df.empty:
        return []

    text = f"""
CSV File: {file_path.name}

Total Rows: {len(df)}
Total Columns: {len(df.columns)}

Columns:
{", ".join(df.columns.astype(str))}

Data Preview:
{df.head(100).to_string(index=False)}
"""

    return [
        Document(
            page_content=text.strip(),
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "csv"
            }
        )
    ]


def load_excel(file_path):
    file_path = Path(file_path)
    docs = []

    excel_file = pd.ExcelFile(file_path)

    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)

        if df.empty:
            continue

        text = f"""
Excel File: {file_path.name}
Sheet Name: {sheet_name}

Total Rows: {len(df)}
Total Columns: {len(df.columns)}

Columns:
{", ".join(df.columns.astype(str))}

Data Preview:
{df.head(100).to_string(index=False)}
"""

        docs.append(
            Document(
                page_content=text.strip(),
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "excel",
                    "sheet": sheet_name
                }
            )
        )

    return docs


def load_single_file(file_path):
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        return load_txt(file_path)

    if suffix == ".md":
        return load_txt(file_path)

    if suffix == ".pdf":
        return load_pdf(file_path)

    if suffix == ".docx":
        return load_docx(file_path)

    if suffix == ".csv":
        return load_csv(file_path)

    if suffix in [".xlsx", ".xls"]:
        return load_excel(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")


# --------------------------------------------------
# Format retrieved documents
# --------------------------------------------------
def format_docs(docs):
    formatted_docs = []

    for doc in docs:
        file_name = doc.metadata.get("file_name", "Unknown file")
        page = doc.metadata.get("page")
        sheet = doc.metadata.get("sheet")

        source_info = f"Source: {file_name}"

        if page:
            source_info += f", Page: {page}"

        if sheet:
            source_info += f", Sheet: {sheet}"

        formatted_docs.append(
            f"{source_info}\n{doc.page_content}"
        )

    return "\n\n---\n\n".join(formatted_docs)


# --------------------------------------------------
# Build Ollama RAG pipeline
# --------------------------------------------------
def build_rag(file_paths):
    """
    Build a local RAG pipeline using Ollama + FAISS.

    Returns:
        rag_chain, stats
    """

    start_time = time.time()

    ollama_base_url = os.getenv(
        "OLLAMA_BASE_URL",
        "http://localhost:11434"
    )

    llm_model_name = os.getenv(
        "OLLAMA_LLM_MODEL",
        "qwen3:4b"
    )

    embedding_model_name = os.getenv(
        "OLLAMA_EMBED_MODEL",
        "mxbai-embed-large:latest"
    )

    # -----------------------------
    # Load all documents
    # -----------------------------
    all_documents = []

    for file_path in file_paths:
        docs = load_single_file(file_path)
        all_documents.extend(docs)

    print("Documents loaded:", len(all_documents))

    if not all_documents:
        raise ValueError(
            "No text was extracted from uploaded files. "
            "Try TXT, text-based PDF, DOCX, CSV, or XLSX. "
            "Scanned PDFs need OCR."
        )

    # -----------------------------
    # Split into chunks
    # -----------------------------
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=120
    )

    chunks = text_splitter.split_documents(all_documents)

    print("Chunks created:", len(chunks))

    if not chunks:
        raise ValueError(
            "Documents were loaded, but no chunks were created."
        )

    # -----------------------------
    # Embeddings with Ollama
    # -----------------------------
    embeddings = OllamaEmbeddings(
        model=embedding_model_name,
        base_url=ollama_base_url
    )

    # -----------------------------
    # FAISS vector store
    # -----------------------------
    vector_store = FAISS.from_documents(
        chunks,
        embeddings
    )

    retriever = vector_store.as_retriever(
        search_kwargs={"k": 4}
    )

    # -----------------------------
    # Local Ollama LLM
    # -----------------------------
    llm = ChatOllama(
        model=llm_model_name,
        base_url=ollama_base_url,
        temperature=0
    )

    # -----------------------------
    # Prompt
    # -----------------------------
    prompt = ChatPromptTemplate.from_template(
        """
You are a helpful local RAG assistant.

Answer the question using only the uploaded document context.

Rules:
- Use only the context.
- Do not use outside knowledge.
- If the answer is not found, say:
  "I could not find this information in the uploaded documents."
- Keep the answer clear and student-friendly.
- Do not show hidden reasoning.
- Do not show thinking tags.

Context:
{context}

Question:
{question}

Answer:
"""
    )

    # -----------------------------
    # RAG chain
    # -----------------------------
    rag_chain = (
        {
            "context": retriever | format_docs,
            "question": RunnablePassthrough()
        }
        | prompt
        | llm
    )

    end_time = time.time()

    stats = {
        "total_files": len(file_paths),
        "total_documents": len(all_documents),
        "total_chunks": len(chunks),
        "embedding_model": embedding_model_name,
        "vector_database": "FAISS",
        "llm_model": llm_model_name,
        "processing_time": round(end_time - start_time, 2),
        "status": "Ready with Ollama"
    }

    return rag_chain, stats